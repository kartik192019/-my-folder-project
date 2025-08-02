from flask import Flask, render_template, request, jsonify, flash, redirect, url_for
import psycopg2
from psycopg2.extras import RealDictCursor
import uuid
import os
from datetime import datetime
import requests
import json
from werkzeug.utils import secure_filename
import PyPDF2
import docx
import chardet
import aspose.words as aw
from supabase import create_client, Client
from config import Config

app = Flask(__name__)
app.secret_key = 'your-secret-key-here'  # Change this to a secure secret key

# Add custom filter for JSON parsing in templates
@app.template_filter('from_json')
def from_json_filter(value):
    """Custom filter to parse JSON in templates"""
    try:
        return json.loads(value)
    except (ValueError, TypeError):
        return value

# Database configuration
DB_CONFIG = Config.get_db_config()

# Supabase configuration
supabase: Client = create_client(Config.SUPABASE_URL, Config.SUPABASE_KEY)

# Backend AI analyzer URL
AI_ANALYZER_URL = 'http://localhost:5001/analyze'

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'doc', 'docx'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Create upload folder if it doesn't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def convert_doc_to_docx(file_obj, filename):
    """Convert .doc file to .docx format in memory"""
    try:
        # Reset file pointer to beginning
        file_obj.seek(0)
        
        # Create a temporary file path for the input
        import tempfile
        import os
        
        # Create temporary files
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_doc:
            temp_doc.write(file_obj.read())
            temp_doc_path = temp_doc.name
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            temp_docx_path = temp_docx.name
        
        try:
            # Load the .doc document
            doc = aw.Document(temp_doc_path)
            
            # Save the document as .docx
            doc.save(temp_docx_path)
            
            # Read the converted .docx file
            with open(temp_docx_path, 'rb') as converted_file:
                converted_content = converted_file.read()
            
            return converted_content
            
        finally:
            # Clean up temporary files
            try:
                os.unlink(temp_doc_path)
                os.unlink(temp_docx_path)
            except:
                pass
                
    except Exception as e:
        print(f"Error converting .doc to .docx: {e}")
        raise Exception(f"Failed to convert .doc file: {str(e)}")

def upload_to_supabase_storage(file, filename):
    """Upload file to Supabase storage and return the public URL"""
    try:
        print(f"Starting Supabase upload for {filename}")
        
        # Check if it's a .doc file and convert to .docx
        file_extension = os.path.splitext(filename)[1].lower()
        original_filename = filename
        
        if file_extension == '.doc':
            print(f"Converting .doc file to .docx format...")
            try:
                # Convert .doc to .docx
                converted_content = convert_doc_to_docx(file, filename)
                
                # Update filename to .docx
                filename = os.path.splitext(filename)[0] + '.docx'
                print(f"Converted {original_filename} to {filename}")
                
                # Use converted content for upload
                file_content = converted_content
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            except Exception as e:
                print(f"Failed to convert .doc to .docx: {e}")
                # Fallback to original file
                file.seek(0)
                file_content = file.read()
                file.seek(0)
                content_type = file.content_type
        else:
            # Read original file content
            file_content = file.read()
            file.seek(0)  # Reset file pointer for later use
            content_type = file.content_type
        
        # Upload to Supabase storage
        bucket_name = Config.SUPABASE_STORAGE_BUCKET
        file_path = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{filename}"
        
        print(f"Uploading to bucket: {bucket_name}, path: {file_path}")
        
        # Upload the file
        result = supabase.storage.from_(bucket_name).upload(
            path=file_path,
            file=file_content,
            file_options={"content-type": content_type}
        )
        
        print(f"Upload result: {result}")
        
        if result:
            # Get the public URL
            public_url = supabase.storage.from_(bucket_name).get_public_url(file_path)
            print(f"Generated public URL: {public_url}")
            return public_url, file_path
        else:
            raise Exception("Failed to upload file to Supabase storage")
            
    except Exception as e:
        print(f"Supabase upload error: {e}")
        # Fallback to local storage if Supabase fails
        print("Falling back to local storage...")
        
        # Handle .doc conversion for local storage too
        if file_extension == '.doc':
            try:
                converted_content = convert_doc_to_docx(file, filename)
                local_filename = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{filename}"
                local_path = os.path.join(app.config['UPLOAD_FOLDER'], local_filename)
                
                # Save converted content
                with open(local_path, 'wb') as f:
                    f.write(converted_content)
                print(f"Saved converted .docx to local path: {local_path}")
                return None, local_filename
            except Exception as conv_e:
                print(f"Failed to convert .doc for local storage: {conv_e}")
                # Fallback to original file
                file.seek(0)
                local_filename = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{original_filename}"
                local_path = os.path.join(app.config['UPLOAD_FOLDER'], local_filename)
                file.save(local_path)
                print(f"Saved original .doc to local path: {local_path}")
                return None, local_filename
        else:
            local_filename = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{filename}"
            local_path = os.path.join(app.config['UPLOAD_FOLDER'], local_filename)
            file.save(local_path)
            print(f"Saved to local path: {local_path}")
            return None, local_filename

def clean_text(text):
    """Clean text to remove null characters and problematic content"""
    if not text:
        return ""
    
    # Remove null characters and other problematic control characters
    text = text.replace('\x00', '')
    text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\r\t')
    
    # Normalize whitespace but preserve line breaks
    lines = text.split('\n')
    cleaned_lines = [' '.join(line.split()) for line in lines]
    text = '\n'.join(line for line in cleaned_lines if line.strip())
    
    return text

def extract_text_from_file_object(file_obj, file_extension):
    """Extract text from file object in memory"""
    try:
        # Reset file pointer to beginning
        file_obj.seek(0)
        
        if file_extension.lower() == '.pdf':
            text = ""
            try:
                pdf_reader = PyPDF2.PdfReader(file_obj)
                print(f"PDF has {len(pdf_reader.pages)} pages")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                            print(f"Extracted {len(page_text)} characters from page {page_num + 1}")
                        else:
                            print(f"No text extracted from page {page_num + 1}")
                    except Exception as e:
                        print(f"Error extracting page {page_num + 1}: {e}")
                        continue
                
                if not text.strip():
                    print("Warning: No text extracted from PDF")
                    return "Error: Could not extract text from PDF. The file might be image-based or corrupted."
                
                return clean_text(text)
            except Exception as e:
                print(f"Error reading PDF: {e}")
                return f"Error: Could not read PDF file: {str(e)}"
        
        elif file_extension.lower() == '.docx':
            try:
                doc = docx.Document(file_obj)
                text = ""
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + "\n"
                return clean_text(text)
            except Exception as e:
                print(f"Error reading Word document: {e}")
                return "Error: Could not read Word document. Please try saving as .txt or copy-paste the content."
        
        elif file_extension.lower() == '.doc':
            try:
                # Convert .doc to .docx first
                print("Converting .doc file to .docx format...")
                converted_content = convert_doc_to_docx(file_obj, "temp.doc")
                
                # Create a BytesIO object from the converted content
                from io import BytesIO
                docx_file_obj = BytesIO(converted_content)
                
                # Now read the converted .docx file
                doc = docx.Document(docx_file_obj)
                text = ""
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + "\n"
                return clean_text(text)
            except Exception as e:
                print(f"Error converting or reading .doc file: {e}")
                return "Error: Could not convert or read .doc file. Please try saving as .docx or .txt format."
        
        elif file_extension.lower() == '.txt':
            # Reset file pointer and read content
            file_obj.seek(0)
            raw_data = file_obj.read()
            
            # Detect encoding
            encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'
            
            # Decode the content
            try:
                text = raw_data.decode(encoding)
                return clean_text(text)
            except UnicodeDecodeError:
                # Fallback to utf-8 with error handling
                text = raw_data.decode('utf-8', errors='ignore')
                return clean_text(text)
        
        else:
            return "Unsupported file format"
    
    except Exception as e:
        print(f"Error extracting text from file object: {e}")
        return f"Error extracting text: {str(e)}"

def extract_text_from_file(file_path, file_extension):
    """Extract text from different file formats with better error handling"""
    try:
        if file_extension.lower() == '.pdf':
            text = ""
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    print(f"PDF has {len(pdf_reader.pages)} pages")
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                                print(f"Extracted {len(page_text)} characters from page {page_num + 1}")
                            else:
                                print(f"No text extracted from page {page_num + 1}")
                        except Exception as e:
                            print(f"Error extracting page {page_num + 1}: {e}")
                            continue
                    
                    if not text.strip():
                        print("Warning: No text extracted from PDF")
                        return "Error: Could not extract text from PDF. The file might be image-based or corrupted."
                    
                    return clean_text(text)
            except Exception as e:
                print(f"Error reading PDF: {e}")
                return f"Error: Could not read PDF file: {str(e)}"
        
        elif file_extension.lower() == '.docx':
            try:
                doc = docx.Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + "\n"
                return clean_text(text)
            except Exception as e:
                print(f"Error reading Word document: {e}")
                return "Error: Could not read Word document. Please try saving as .txt or copy-paste the content."
        
        elif file_extension.lower() == '.doc':
            try:
                # Convert .doc to .docx first
                print("Converting .doc file to .docx format...")
                
                # Load the .doc document
                doc = aw.Document(file_path)
                
                # Create a temporary file for the converted .docx
                import tempfile
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
                    temp_docx_path = temp_docx.name
                
                try:
                    # Save the document as .docx
                    doc.save(temp_docx_path)
                    
                    # Now read the converted .docx file
                    docx_doc = docx.Document(temp_docx_path)
                    text = ""
                    for paragraph in docx_doc.paragraphs:
                        if paragraph.text.strip():
                            text += paragraph.text + "\n"
                    return clean_text(text)
                    
                finally:
                    # Clean up temporary file
                    try:
                        import os
                        os.unlink(temp_docx_path)
                    except:
                        pass
                        
            except Exception as e:
                print(f"Error converting or reading .doc file: {e}")
                return "Error: Could not convert or read .doc file. Please try saving as .docx or .txt format."
        
        elif file_extension.lower() == '.txt':
            # Detect encoding first
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'
            
            # Read with detected encoding
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                return clean_text(text)
            except UnicodeDecodeError:
                # Fallback to utf-8 with error handling
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    text = file.read()
                return clean_text(text)
        
        else:
            return "Unsupported file format"
    
    except Exception as e:
        print(f"Error extracting text from {file_path}: {e}")
        return f"Error extracting text: {str(e)}"

def get_db_connection():
    """Get database connection with better error handling"""
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        return conn
    except Exception as e:
        print(f"Database connection error: {e}")
        return None

@app.route('/')
def index():
    """Main page with job description upload form"""
    return redirect(url_for('render_upload_page'))

@app.route('/upload', methods=['GET'])
def render_upload_page():
    """Render the job description upload form (GET request)"""
    return render_template('upload.html')

@app.route('/upload', methods=['POST'])
def upload_job_description():
    """Handle job description upload and processing"""
    try:
        if 'file' not in request.files and 'job_description' not in request.form:
            flash('No file or job description provided')
            return redirect(request.url)

        company_id = request.form.get('company_id')
        user_id = request.form.get('user_id')
        title = request.form.get('title')
        
        # Clean and validate UUID fields
        if user_id and (user_id.strip() == '' or user_id == 'null' or user_id == 'undefined'):
            user_id = None
        if company_id and (company_id.strip() == '' or company_id == 'null' or company_id == 'undefined'):
            company_id = None
        
        job_description_text = ""
        jd_file = None
        
        # Handle file upload
        if 'file' in request.files and request.files['file'].filename != '':
            file = request.files['file']
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                
                try:
                    # Upload to Supabase storage
                    print(f"Uploading file: {filename}")
                    supabase_url, file_path = upload_to_supabase_storage(file, filename)
                    jd_file = file_path
                    print(f"Upload result - URL: {supabase_url}, Path: {file_path}")
                    
                    # Extract text based on file type
                    file_extension = os.path.splitext(filename)[1]
                    
                    # If Supabase upload was successful, use file object
                    if supabase_url:
                        print("Extracting text from file object (Supabase upload successful)")
                        job_description_text = extract_text_from_file_object(file, file_extension)
                    else:
                        # Fallback to local file processing
                        print("Extracting text from local file (Supabase upload failed)")
                        local_file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_path)
                        job_description_text = extract_text_from_file(local_file_path, file_extension)
                    
                    if job_description_text.startswith("Error"):
                        flash(f'File processing error: {job_description_text}')
                        return redirect(request.url)
                        
                except Exception as e:
                    flash(f'Error saving file: {str(e)}')
                    return redirect(request.url)
            else:
                flash('Invalid file type. Allowed types: TXT, PDF, DOC, DOCX')
                return redirect(request.url)
        
        # Handle text input
        elif 'job_description' in request.form:
            raw_text = request.form['job_description']
            job_description_text = clean_text(raw_text)
        
        if not job_description_text or not job_description_text.strip():
            flash('Job description is empty or could not be processed')
            return redirect(request.url)

        # Save to database
        conn = get_db_connection()
        if not conn:
            # Try to proceed without database for testing
            flash('Database connection failed, but will attempt analysis anyway')
            
            # Create a mock JD ID for analysis
            jd_id = str(uuid.uuid4())
            jd_file_ref = jd_file if jd_file else jd_id
            
            # Proceed with analysis
            backend_payload = {
                'jd_id': jd_id,
                'jd_file': jd_file_ref,
                'job_description': job_description_text,
                'user_id': user_id
            }
            
            try:
                response = requests.post(AI_ANALYZER_URL, json=backend_payload, timeout=60)
                
                if response.status_code == 200:
                    return jsonify({
                        'status': 'success',
                        'message': 'Job description analyzed successfully (no database save)',
                        'jd_id': jd_id,
                        'analysis_result': response.json()
                    })
                else:
                    return jsonify({
                        'status': 'error',
                        'message': f'Analysis failed: {response.text}',
                        'jd_id': jd_id
                    })
                    
            except requests.exceptions.RequestException as e:
                return jsonify({
                    'status': 'error',
                    'message': f'Backend connection failed: {str(e)}'
                })

        try:
            cur = conn.cursor()
            
            # For file uploads, use the Supabase URL as file_id (only if upload was successful)
            file_id = supabase_url if ('file' in request.files and request.files['file'].filename != '' and supabase_url) else None
            
            insert_query = """
                INSERT INTO job_descriptions (company_id, user_id, title, jd_file, file_id)
                VALUES (%s, %s, %s, %s, %s)
                RETURNING jd_id, jd_file, file_id
            """
            
            # Clean all inputs before database insertion
            # Handle UUID conversion safely
            company_uuid = None
            user_uuid = None
            
            if company_id:
                try:
                    company_uuid = uuid.UUID(company_id)
                except (ValueError, TypeError):
                    print(f"Invalid company_id format: {company_id}")
                    company_uuid = None
            
            if user_id:
                try:
                    user_uuid = uuid.UUID(user_id)
                except (ValueError, TypeError):
                    print(f"Invalid user_id format: {user_id}")
                    user_uuid = None
            
            cur.execute(insert_query, (
                company_uuid,
                user_uuid,
                clean_text(title) if title else None,
                jd_file,
                file_id
            ))
            
            result = cur.fetchone()
            jd_id = result[0]
            jd_file_ref = result[1] if result[1] else str(jd_id)
            file_id_ref = result[2] if result[2] else str(jd_id)
            
            conn.commit()
            cur.close()
            conn.close()

            # Call backend for analysis
            # Use the file_id (Supabase URL) as referenced_jd for the analysis
            referenced_jd_url = file_id_ref if file_id_ref else str(jd_id)
            backend_payload = {
                'jd_id': str(jd_id),
                'jd_file': referenced_jd_url,  # Use Supabase URL as referenced_jd
                'job_description': job_description_text,
                'user_id': user_id
            }

            try:
                print(f"Sending analysis request to: {AI_ANALYZER_URL}")
                print(f"Payload: {json.dumps(backend_payload, indent=2)}")
                
                response = requests.post(AI_ANALYZER_URL, json=backend_payload, timeout=60)
                
                print(f"Response status: {response.status_code}")
                print(f"Response text: {response.text}")
                
                if response.status_code == 200:
                    flash('Job description uploaded and analyzed successfully!')
                    return jsonify({
                        'status': 'success',
                        'message': 'Job description processed successfully',
                        'jd_id': str(jd_id),
                        'analysis_result': response.json()
                    })
                else:
                    error_msg = f'Analysis failed: {response.text}'
                    print(f"Analysis error: {error_msg}")
                    flash(error_msg)
                    return jsonify({
                        'status': 'partial_success',
                        'message': 'Job description saved but analysis failed',
                        'jd_id': str(jd_id),
                        'error_details': response.text
                    })
                    
            except requests.exceptions.RequestException as e:
                error_msg = f'Backend connection failed: {str(e)}'
                print(f"Connection error: {error_msg}")
                flash(error_msg)
                return jsonify({
                    'status': 'partial_success',
                    'message': 'Job description saved but analysis failed',
                    'jd_id': str(jd_id),
                    'error_details': str(e)
                })

        except Exception as e:
            if conn:
                conn.rollback()
                cur.close()
                conn.close()
            flash(f'Database error: {str(e)}')
            return redirect(request.url)

    except Exception as e:
        flash(f'Upload error: {str(e)}')
        return redirect(request.url)

@app.route('/job_descriptions')
def list_job_descriptions():
    """List all job descriptions"""
    conn = get_db_connection()
    if not conn:
        flash('Database connection failed')
        return render_template('list.html', job_descriptions=[])

    try:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute("""
            SELECT jd_id, title, company_id, user_id, created_at, updated_at, file_id, jd_file
            FROM job_descriptions
            ORDER BY created_at DESC
            LIMIT 100
        """)
        job_descriptions = cur.fetchall()
        cur.close()
        conn.close()
        
        return render_template('list.html', job_descriptions=job_descriptions)
        
    except Exception as e:
        flash(f'Error fetching job descriptions: {str(e)}')
        return render_template('list.html', job_descriptions=[])

@app.route('/job_description/<jd_id>')
def view_job_description(jd_id):
    """View specific job description and its analysis"""
    conn = get_db_connection()
    if not conn:
        flash('Database connection failed')
        return redirect(url_for('list_job_descriptions'))

    try:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        cur.execute("""SELECT * FROM job_descriptions WHERE jd_id = %s""", (jd_id,))
        jd = cur.fetchone()
        
        if not jd:
            flash('Job description not found')
            return redirect(url_for('list_job_descriptions'))
        
        cur.execute("""
            SELECT * FROM resolved_jd 
            WHERE referenced_jd = %s OR referenced_jd = %s OR referenced_jd = %s
            ORDER BY created_at DESC
        """, (jd['file_id'], jd['jd_file'], str(jd['jd_id'])))
        
        resolved_data = cur.fetchall()
        cur.close()
        conn.close()
        
        return render_template('view.html', job_description=jd, resolved_data=resolved_data)
        
    except Exception as e:
        flash(f'Error fetching job description: {str(e)}')
        return redirect(url_for('list_job_descriptions'))

@app.route('/test')
def test_text_cleaning():
    """Test endpoint to check text cleaning"""
    test_text = "Hello\x00World\x01Test\x02"
    cleaned = clean_text(test_text)
    return jsonify({
        'original': repr(test_text),
        'cleaned': repr(cleaned),
        'original_length': len(test_text),
        'cleaned_length': len(cleaned)
    })

if __name__ == '__main__':
    print("Starting JD Upload Application on http://127.0.0.1:5004")
    app.run(debug=True, host='127.0.0.1', port=5004)